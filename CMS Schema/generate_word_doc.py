#!/usr/bin/env python3
"""
Generates a Word document with ALL fields for ALL recipes from the updated CSV.
"""

import csv
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

INPUT_FILE = os.path.join(os.path.dirname(__file__), "Recipe_schema_UPDATED.csv")
OUTPUT_FILE = os.path.join(os.path.dirname(__file__), "Recipe_CMS_Updates.docx")

# Human-readable labels for CSV columns
FIELD_LABELS = {
    "Slug": "Slug",
    "Name": "Recipe Name",
    "Created Date": "Created Date",
    "Updated Date": "Updated Date",
    "Short Description": "Short Description",
    "Featured Image": "Featured Image",
    "Featured Image:alt": "Featured Image Alt Text",
    "Image": "Image",
    "Image:alt": "Image Alt Text",
    "Image 2": "Image 2",
    "Image 2:alt": "Image 2 Alt Text",
    "Image 3": "Image 3",
    "Image 3:alt": "Image 3 Alt Text",
    "Image 4": "Image 4",
    "Image 4:alt": "Image 4 Alt Text",
    "Image 5": "Image 5",
    "Image 5:alt": "Image 5 Alt Text",
    "Image 6": "Image 6",
    "Image 6:alt": "Image 6 Alt Text",
    "Introduction": "Introduction",
    "Why You'll Love This": "Why You'll Love This",
    "Tips": "Tips",
    "Ingredients Overview": "Ingredients Overview",
    "Instructions Overview": "Instructions Overview",
    "FAQ": "FAQ",
    "author": "Author",
    "Prep Time": "Prep Time",
    "Cook Time": "Cook Time",
    "Total Time": "Total Time",
    "Servings": "Servings",
    "Calories": "Calories",
    "Category": "Category",
    "Cuisine": "Cuisine",
    "Recipe Card Ingredients": "Recipe Card Ingredients",
    "Recipe card Instructions": "Recipe Card Instructions",
    "Nutrition": "Nutrition",
    "Source URL": "Source URL",
    "Diet Types": "Diet Types",
    "Categories": "Categories",
    "Cuisines": "Cuisines",
    "Related Recipes": "Related Recipes",
    "Diet": "Diet",
    "Features": "Features",
}

# Skip image URL fields (long URLs that aren't useful in a Word doc)
SKIP_IF_EMPTY = {
    "Featured Image", "Featured Image:alt",
    "Image", "Image:alt",
    "Image 2", "Image 2:alt",
    "Image 3", "Image 3:alt",
    "Image 4", "Image 4:alt",
    "Image 5", "Image 5:alt",
    "Image 6", "Image 6:alt",
}


def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def main():
    with open(INPUT_FILE, "r", encoding="utf-8") as f:
        reader = csv.DictReader(f)
        fieldnames = list(reader.fieldnames)
        rows = list(reader)

    print(f"Read {len(rows)} recipes with {len(fieldnames)} fields each")

    doc = Document()

    # Adjust default style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    # Title
    title = doc.add_heading("Holistic Bravo - Complete Recipe Data", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Intro
    intro = doc.add_paragraph()
    intro.add_run(
        f"This document contains all CMS fields for all {len(rows)} recipes.\n"
        "Generated from Recipe_schema_UPDATED.csv\n\n"
        "Image fields are omitted when empty to save space."
    ).font.size = Pt(10)

    # Table of contents style - list all recipes
    doc.add_heading("Recipe Index", level=1)
    for i, row in enumerate(rows):
        p = doc.add_paragraph(f"{i + 1}. {row.get('Name', '')}", style="List Number")
        p.paragraph_format.space_after = Pt(2)

    doc.add_page_break()

    for i, row in enumerate(rows):
        name = row.get("Name", "")

        # Recipe heading
        heading = doc.add_heading(f"{i + 1}. {name}", level=1)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)

        # Create table with all fields
        table = doc.add_table(rows=0, cols=2)
        table.style = "Table Grid"
        table.autofit = True

        for col_name in fieldnames:
            value = row.get(col_name, "").strip()

            # Skip empty image fields
            if col_name in SKIP_IF_EMPTY and not value:
                continue

            label = FIELD_LABELS.get(col_name, col_name)

            tbl_row = table.add_row()

            # Field name cell
            cell0 = tbl_row.cells[0]
            cell0.width = Inches(1.8)
            p0 = cell0.paragraphs[0]
            run0 = p0.add_run(label)
            run0.bold = True
            run0.font.size = Pt(9)
            run0.font.name = "Calibri"
            set_cell_shading(cell0, "F0E6FF")

            # Field value cell
            cell1 = tbl_row.cells[1]
            cell1.width = Inches(5)
            p1 = cell1.paragraphs[0]

            if value:
                # For long text, keep font smaller
                font_size = Pt(9) if len(value) > 200 else Pt(10)
                run1 = p1.add_run(value)
                run1.font.size = font_size
                run1.font.name = "Calibri"
            else:
                run1 = p1.add_run("(empty)")
                run1.font.size = Pt(9)
                run1.font.name = "Calibri"
                run1.font.color.rgb = RGBColor(153, 153, 153)
                run1.italic = True

        # Page break after each recipe except the last
        if i < len(rows) - 1:
            doc.add_page_break()

    doc.save(OUTPUT_FILE)
    print(f"Word document saved to: {OUTPUT_FILE}")


if __name__ == "__main__":
    main()
